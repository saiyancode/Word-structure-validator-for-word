from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from nltk import ngrams
import nltk
import re

""" Objective highlight words in a word doc that do not match the structure which is set"""

#TODO conserve formatting of a rebuilt paragraph


# Open our document & return a list of paragraphs to analyse
document = Document('test.docx')
paragraph_loop = document.paragraphs

# Open the text file we've populated with words we need to check
words_to_check =  [line.rstrip('\n') for line in open('words_to_check.txt')]

def rebuild_para(para, text,mismatched):
    # Split the original paragraph up with all the values in mismatched to account for multiple matches per paragraph
    try:
        para_split = re.split("{}{}{}".format("(","|".join(list(mismatched)),")"),text, flags=re.IGNORECASE) # Case insensitive regex split the conserves our split terms
        p2 = para.insert_paragraph_before() # Inserts a new blank para to build into
        # Builds a new paragraph that has the words highlighted
        for val in para_split:
            if val.lower() in mismatched:
                p2.add_run(val).font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                p2.add_run(val)

        # Maintain bold elements
        # TODO

        # Maintain italic elements
        #TODO

        # We need to remove the original paragraph as we've created our highlighted
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = Nones

    except Exception as e:
        print(e)
        pass

def check_text(para):
    # Split the paragraph into the corresponding lengths of the words in the checklist
    results = set()
    for phrase in words_to_check:
        phrase_length = len(phrase.split())
        # Split the paragraph into ngrams of the length of our check word
        tokenize = nltk.word_tokenize(phrase)
        token_list = [ngram for ngram in ngrams(tokenize,phrase_length)]

        # tokenise the paragaph in the same way

        tokenize_para = nltk.word_tokenize(para)
        para_list = [ngram for ngram in ngrams(tokenize_para,phrase_length)]

        # Loop over the two lists checking each ngram in the paragraph against the list of words to check if we find a mismatch add to the results

        for ngram in para_list:
            for ngram1 in token_list:
                if " ".join(ngram).lower() == " ".join(ngram1).lower():
                    if " ".join(ngram) != " ".join(ngram1):
                        results.add(" ".join(ngram).lower())

    return results

if __name__ == "__main__":
    # Loop over all the content in our word doc
    for count, para in enumerate(paragraph_loop):
        text = para.text
        mismatched = check_text(text)
        rebuild_para(para, text, mismatched)

# Save the outfile as a new name
document.save('test-new.docx')